# -*- coding: utf-8 -*-
# -----------------------------------------------------------------------------
# Name:         urls.py (App Level)
# Purpose:      URL routing for the Invoice application's web pages.
#
# Author:       AnikRoy
# GitHub:       https://github.com/aroyslipk
#
# Created:      2025-06-20
# Copyright:    (c) AnikRoy 2025
# Licence:      Proprietary
# -----------------------------------------------------------------------------



from datetime import datetime
from pathlib import Path
from copy import copy
from io import BytesIO
from collections import Counter

from django.conf import settings
from collections import defaultdict
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import LoginView
from django.contrib.auth import login
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib.auth import login, logout
from django.contrib.auth.forms import AuthenticationForm
from django.http import JsonResponse
import user_agents
from django.core.mail import send_mail
from django.core.files.base import ContentFile
from django.core.paginator import Paginator
from django.db.models import Q, Count, Sum
from django.db.models.functions import ExtractWeekDay
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.dateparse import parse_date
from django.http import JsonResponse
from decimal import Decimal
import json
from django.db.models import Sum, Count
from django.utils import timezone

from num2words import num2words
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Border, Side, Alignment
from rest_framework import generics, permissions
from rest_framework.response import Response
from rest_framework_simplejwt.authentication import JWTAuthentication

from .models import User, ClientProject, Category, WorkEntry, Invoice, UserLoginHistory

from .forms import (
    AdminUserCreationForm,
    ClientProjectForm,
    CategoryForm,
    WorkEntryForm,
    AdminSlotForm,
    UserRoleUpdateForm,
)
from .serializers import (
    ClientProjectSerializer,
    RegisterSerializer,
    UserSerializer,
    WorkDashboardSerializer,
    WorkEntrySerializer,
    CategorySerializer,
)

def home_view(request):
    return render(request, 'home.html')


@login_required
def dashboard_template_view(request):
    user = request.user
    if user.role == 'user':
        return redirect('my_work_entries')

    projects_qs = ClientProject.objects.all()
    projects_qs = ClientProject.objects.all()
    users_qs = User.objects.filter(role='user')
    entries_qs = WorkEntry.objects.select_related('user', 'project', 'category')

    if user.role == 'admin':
        projects_qs = projects_qs.filter(managed_by=user)
        users_qs = users_qs.filter(created_by=user)
        entries_qs = entries_qs.filter(project__managed_by=user)

    selected_project_id = request.GET.get('project')
    selected_user_id = request.GET.get('user')
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')

    if selected_project_id:
        entries_qs = entries_qs.filter(project_id=selected_project_id)
    if selected_user_id:
        entries_qs = entries_qs.filter(user_id=selected_user_id)
    if start_date:
        entries_qs = entries_qs.filter(date__gte=start_date)
    if end_date:
        entries_qs = entries_qs.filter(date__lte=end_date)

    total_projects = projects_qs.count()
    total_team_members = users_qs.count()
    current_month_entries = entries_qs.filter(
        date__year=timezone.now().year,
        date__month=timezone.now().month
    ).count()

    busiest_day_query = entries_qs.annotate(
        weekday=ExtractWeekDay('date')
    ).values('weekday').annotate(count=Count('id')).order_by('-count').first()

    weekday_map = {1: 'Sunday', 2: 'Monday', 3: 'Tuesday', 4: 'Wednesday', 5: 'Thursday', 6: 'Friday', 7: 'Saturday'}
    busiest_day = weekday_map.get(busiest_day_query['weekday'], 'N/A') if busiest_day_query else 'N/A'

    most_productive_user_query = entries_qs.values('user__username').annotate(
        total_quantity=Sum('quantity')
    ).order_by('-total_quantity').first()
    most_productive_user = most_productive_user_query['user__username'] if most_productive_user_query else 'N/A'

    category_data = entries_qs.values('category__name').annotate(count=Count('id')).order_by('-count')
    pie_chart_labels = [item['category__name'] or 'Uncategorized' for item in category_data]
    pie_chart_data = [item['count'] for item in category_data]

    monthly_data = (
        entries_qs.values('date__year', 'date__month')
        .annotate(count=Count('id'))
        .order_by('date__year', 'date__month')
    )
    bar_chart_labels = [f"{item['date__year']}-{item['date__month']:02d}" for item in monthly_data]
    bar_chart_data = [item['count'] for item in monthly_data]

    total_invoiced_amount_query = Invoice.objects.filter(project__in=projects_qs).aggregate(total=Sum('total_amount'))
    total_invoiced_amount = total_invoiced_amount_query['total'] or 0.00

    revenue_month_year = request.GET.get('revenue_month_year')
    selected_year, selected_month = None, None

    if revenue_month_year:
        year_month = revenue_month_year.split('-')
        if len(year_month) == 2:
            selected_year = year_month[0]
            selected_month = year_month[1]

    month_names = {
        "1": "January", "2": "February", "3": "March", "4": "April",
        "5": "May", "6": "June", "7": "July", "8": "August",
        "9": "September", "10": "October", "11": "November", "12": "December"
    }

    available_revenue_dates = WorkEntry.objects.dates('date', 'month', order='DESC')

    monthly_revenue = Decimal('0.00')
    if selected_year and selected_month:
        revenue_entries = entries_qs.filter(
            date__year=selected_year,
            date__month=selected_month
        )
        for entry in revenue_entries:
            rate = entry.category.rate if entry.category else Decimal('0.00')
            quantity = entry.quantity or 0
            monthly_revenue += (Decimal(quantity) * rate)

    page_obj = None
    if entries_qs.exists():
        paginator = Paginator(entries_qs, 18)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

    current_month = datetime.today().strftime('%Y-%m')
    context = {
        "total_projects": total_projects,
        "total_team_members": total_team_members,
        "current_month_entries": current_month_entries,
        'current_month': current_month,
        "busiest_day": busiest_day,
        "most_productive_user": most_productive_user,
        'entries': entries_qs.order_by('-date'),
        "pie_chart_labels": pie_chart_labels,
        "pie_chart_data": pie_chart_data,
        "bar_chart_labels": bar_chart_labels,
        "bar_chart_data": bar_chart_data,
        'all_projects': projects_qs.order_by('name'),
        'all_users': users_qs.order_by('username'),
        "total_invoiced_amount": total_invoiced_amount,
        "monthly_revenue": monthly_revenue,
        "month_names": month_names,
        "available_revenue_dates": available_revenue_dates,
        "selected_project_id": selected_project_id,
        "selected_user_id": selected_user_id,
        "start_date": start_date,
        "end_date": end_date,
        "revenue_month_year": revenue_month_year,
        'projects': projects_qs,
        "users": users_qs,
        "entries": entries_qs,
        "page_obj": page_obj,
    }
    return render(request, "dashboard.html", context)


@login_required
def admin_panel_view(request):
    if request.user.role != 'super_admin':
        return render(request, 'unauthorized.html')

    context = {
        'work_entries': WorkEntry.objects.all().order_by('-date'),
        'categories': Category.objects.all(),
        'projects': ClientProject.objects.all(),
        'users': User.objects.all(),
    }
    return render(request, 'admin_panel.html', context)


@login_required
def work_entry_form_view(request):
    print(f"--- DEBUGGING --- User: {request.user.username}, Created by: {request.user.created_by}")
    if request.user.role != 'user':
        return render(request, 'unauthorized.html')

    admin_of_user = request.user.created_by

    if request.method == 'POST':
        form = WorkEntryForm(request.POST, admin_manager=admin_of_user)
        if form.is_valid():
            work_entry = form.save(commit=False)
            work_entry.user = request.user
            work_entry.save()
            return redirect('my_work_entries')
    else:
        form = WorkEntryForm(admin_manager=admin_of_user)

    return render(request, 'work_entry_form.html', {'form': form})


@login_required
def my_work_entries_view(request):
    if request.user.role != 'user':
        return render(request, 'unauthorized.html')

    user = request.user
    
    # Start with base queryset
    entries_queryset = WorkEntry.objects.filter(user=user).select_related('project', 'category')
    
    # Get filter parameters
    project_filter = request.GET.get('project', '').strip()
    query = request.GET.get('query', '').strip()
    start_date = request.GET.get('start_date', '').strip()
    end_date = request.GET.get('end_date', '').strip()
    
    # DEBUG: Print initial state
    print(f"DEBUG: Initial entries count: {entries_queryset.count()}")
    print(f"DEBUG: Project filter: '{project_filter}'")
    
    # Apply filters sequentially
    if project_filter:
        before_count = entries_queryset.count()
        entries_queryset = entries_queryset.filter(project__name__iexact=project_filter)
        after_count = entries_queryset.count()
        print(f"DEBUG: After project filter - Before: {before_count}, After: {after_count}")
        print(f"DEBUG: Available project names: {list(WorkEntry.objects.filter(user=user).values_list('project__name', flat=True).distinct())}")
    
    if query:
        entries_queryset = entries_queryset.filter(
            Q(folder_name__icontains=query) |
            Q(category__name__icontains=query)
        )
    
    if start_date:
        try:
            start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date()
            entries_queryset = entries_queryset.filter(date__gte=start_date_obj)
        except ValueError:
            pass
            
    if end_date:
        try:
            end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date()
            entries_queryset = entries_queryset.filter(date__lte=end_date_obj)
        except ValueError:
            pass
    
    # Apply final ordering
    entries_queryset = entries_queryset.order_by('-date')
    
    # DEBUG: Print final state
    print(f"DEBUG: Final entries count: {entries_queryset.count()}")
    print(f"DEBUG: Final entries projects: {[e.project.name for e in entries_queryset[:5]]}")
    
    # Summary stats based on current month and filtered entries
    now = timezone.now()
    current_month_filtered = entries_queryset.filter(date__year=now.year, date__month=now.month)
    total_entries_month = current_month_filtered.count()
    total_quantity_month = entries_queryset.aggregate(total=Sum('quantity'))['total'] or 0

    # Most frequent project from filtered entries
    project_counts = entries_queryset.values('project__name').annotate(
        count=Count('id')
    ).order_by('-count').first()
    most_frequent_project = project_counts['project__name'] if project_counts else 'N/A'

    # Unique projects for dropdown (all user's projects)
    unique_projects = WorkEntry.objects.filter(user=user).values_list('project__name', flat=True).distinct().order_by('project__name')
    
    # Calendar events from filtered entries only
    calendar_events = []
    for entry_date in entries_queryset.values_list('date', flat=True).distinct():
        if entry_date:  # Make sure date is not None
            calendar_events.append({
                'title': 'Work Submitted',
                'start': entry_date.strftime('%Y-%m-%d'),
                'allDay': True
            })

    context = {
        'entries': entries_queryset,
        'total_entries_month': total_entries_month,
        'total_quantity_month': total_quantity_month,
        'most_frequent_project': most_frequent_project,
        'calendar_events_json': json.dumps(calendar_events),
        'start_date': start_date,
        'end_date': end_date,
        'query': query,
        'project_filter': project_filter,
        'unique_projects': unique_projects,
    }

    # Handle AJAX requests
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        entries_html = render_to_string('partials/_work_entries_list.html', {'entries': entries_queryset}, request=request)
        return JsonResponse({
            'entries_html': entries_html,
            'summary': {
                'total_entries_month': total_entries_month,
                'total_quantity_month': total_quantity_month,
                'most_frequent_project': most_frequent_project
            }
        })

    return render(request, 'my_work_entries.html', context)


@login_required
def create_slots_view(request):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        form = AdminSlotForm(request.POST)
        if form.is_valid():
            project = form.cleaned_data['project']
            slot_count = form.cleaned_data['slot_count']
            for _ in range(slot_count):
                WorkEntry.objects.create(project=project, is_slot=True)
            messages.success(request, f"{slot_count} empty slots created for project '{project.name}'")
            return redirect('manage_projects')
    else:
        form = AdminSlotForm()

    return render(request, 'create_slots.html', {'form': form})


@login_required
def fill_slots_view(request):
    if request.user.role != 'user':
        return render(request, 'unauthorized.html')

    user = request.user
    empty_slots = WorkEntry.objects.filter(user__isnull=True, is_slot=True, project__in=user.created_by.managed_projects.all())

    if request.method == 'POST':
        for slot in empty_slots:
            category = request.POST.get(f'category_{slot.id}')
            quantity = request.POST.get(f'quantity_{slot.id}')
            if category and quantity:
                slot.user = user
                slot.category = category
                slot.quantity = quantity
                slot.date = timezone.now()
                slot.save()
        messages.success(request, "Slots filled successfully.")
        return redirect('my_work_entries')

    return render(request, 'fill_slots.html', {'slots': empty_slots})

@login_required
def submit_work_view(request):
    print("=== Starting submit_work_view ===")
    print(f"User: {request.user.username}")

    admin_of_user = request.user.created_by
    if not admin_of_user:
        projects = ClientProject.objects.none()
        print("No admin found for user")
    else:
        projects = ClientProject.objects.filter(managed_by=admin_of_user)
        print(f"Found projects: {[p.name for p in projects]}")

    if request.method == 'POST':
        print("Processing POST request")
        project_id = request.POST.get('project')
        category_ids = request.POST.getlist('category[]')
        folder_names = request.POST.getlist('folder_name[]')
        quantities = request.POST.getlist('quantity[]')
        dates = request.POST.getlist('date[]')

        print(f"Received data - Project ID: {project_id}")
        print(f"Categories: {category_ids}")
        print(f"Folders: {folder_names}")
        print(f"Quantities: {quantities}")
        print(f"Dates: {dates}")

        if not project_id:
            messages.error(request, "Please select a project.")
            return redirect('submit_work')

        project = get_object_or_404(ClientProject, id=project_id)
        entries_created = []

        for cat_id, folder, qty, date_str in zip(category_ids, folder_names, quantities, dates):
            if folder.strip() and qty and date_str and cat_id:
                try:
                    category_instance = Category.objects.get(id=cat_id)
                    # Convert date string to timezone-aware datetime
                    date_obj = timezone.datetime.strptime(date_str, '%Y-%m-%d')
                    local_tz = timezone.get_current_timezone()
                    aware_date = timezone.make_aware(date_obj, local_tz)
                    
                    entry = WorkEntry.objects.create(
                        user=request.user,
                        project=project,
                        folder_name=folder.strip(),
                        category=category_instance,
                        quantity=int(qty),
                        date=aware_date
                    )
                    if entry.id:
                        entries_created.append(entry)
                except (Category.DoesNotExist, ValueError) as e:
                    print(f"Error creating entry: {str(e)}")
                    messages.error(request, f"Error creating entry: {str(e)}")
                    continue

        if entries_created:
            messages.success(request, f"{len(entries_created)} work entries submitted successfully!")
        else:
            messages.warning(request, "No work entries were created. Please check your input.")

        return redirect('my_work_entries')

    context = {
        'projects': projects,
        'categories': Category.objects.filter(managed_by=admin_of_user) if admin_of_user else Category.objects.none(),
        'today': timezone.now()
    }
    return render(request, 'submit_work.html', context)

@login_required
def manage_prices_view(request, project_id):
    project = get_object_or_404(ClientProject, id=project_id)
    user = request.user

    if user.role != 'super_admin' and project.managed_by != user:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        if 'add_category' in request.POST:
            form = CategoryForm(request.POST)
            if form.is_valid():
                category = form.save(commit=False)
                category.project = project
                category.managed_by = user
                category.save()
                messages.success(request, f"Category '{category.name}' added to project '{project.name}'.")

        elif 'update_prices' in request.POST:
            for key, value in request.POST.items():
                if key.startswith('rate-'):
                    try:
                        category_id = int(key.split('-')[1])
                        category = Category.objects.get(id=category_id, project=project, managed_by=user)
                        category.rate = value
                        category.save()
                    except (ValueError, Category.DoesNotExist):
                        continue
            messages.success(request, f"Prices for project '{project.name}' updated successfully.")

        return redirect('manage_prices', project_id=project.id)

    form = CategoryForm()
    categories = project.categories.all().order_by('name')

    context = {
        'form': form,
        'project': project,
        'categories': categories,
    }
    return render(request, 'manage_prices.html', context)

@login_required
def set_price_view(request, entry_id):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        work_entry = get_object_or_404(WorkEntry, id=entry_id)
        price = request.POST.get('price')

        if price and price.strip():
            try:
                work_entry.unit_price = float(price)
                work_entry.save()
                messages.success(request, f"Price set for folder '{work_entry.category}'.")
            except ValueError:
                messages.error(request, "Invalid price format.")
        else:
            work_entry.unit_price = None
            work_entry.save()
            messages.warning(request, f"Price removed for folder '{work_entry.category}'.")

    return redirect('manage_projects')

@login_required
def price_edit_view(request, id):
    price = get_object_or_404(Category, id=id)
    user = request.user
    if user.role != 'super_admin' and price.managed_by != user:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        form = CategoryForm(request.POST, instance=price)
        if form.is_valid():
            form.save()
            return redirect('manage_prices', project_id=price.project.id)
    else:
        form = CategoryForm(instance=price)
    return render(request, 'price_edit.html', {'form': form, 'category': price})


@login_required
def delete_price_view(request, id):
    user = request.user
    price = get_object_or_404(Category, id=id)

    if user.role != 'super_admin' and price.managed_by != user:
        return render(request, 'unauthorized.html')

    project_id_to_redirect = price.project.id
    category_name = price.name
    price.delete()
    messages.success(request, f"Category '{category_name}' has been deleted.")
    return redirect('manage_prices', project_id=project_id_to_redirect)


@login_required
def save_all_prices_view(request, project_id):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        entries_to_update = []
        for key, value in request.POST.items():
            if key.startswith('price-'):
                try:
                    entry_id = int(key.split('-')[1])
                    entry = WorkEntry.objects.get(id=entry_id, project_id=project_id)
                    entry.unit_price = float(value) if value else None
                    entries_to_update.append(entry)
                except (ValueError, WorkEntry.DoesNotExist):
                    continue

        if entries_to_update:
            WorkEntry.objects.bulk_update(entries_to_update, ['unit_price'])
            messages.success(request, "All prices have been updated successfully.")

    return redirect('manage_projects')

@login_required
def export_page_view(request):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')
    return render(request, 'export_page.html')

@login_required
def my_team_view(request):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        form = AdminUserCreationForm(request.POST, user=request.user)
        if form.is_valid():
            password = form.cleaned_data.get('password2')
            new_user = form.save(commit=False)
            if request.user.role == 'super_admin':
                new_user.role = form.cleaned_data.get('role', 'user')  # Allow super_admin to set role
            else:
                new_user.role = 'user'  # Admins can only create regular users
            new_user.created_by = request.user
            new_user.save()

            try:
                mail_subject = 'Welcome to InvoiceApp!'
                message = render_to_string('welcome_email.html', {
                    'user': new_user,
                    'admin_name': request.user.username,
                    'password': password,
                })

                send_mail(
                    subject=mail_subject,
                    message=message,
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[new_user.email],
                    html_message=message,
                    fail_silently=False
                )
                messages.success(request, f"User '{new_user.username}' was created successfully! A welcome email has been sent.")
            except Exception as e:
                messages.warning(request, f"User '{new_user.username}' created, but failed to send welcome email. Error: {e}")

            return redirect('my_team')
    else:
        form = AdminUserCreationForm()

    # For super_admin: show all users except themselves
    if request.user.role == 'super_admin':
        managed_users = User.objects.exclude(id=request.user.id).select_related('created_by')
    # For admin: only show their managed users
    else:
        managed_users = User.objects.filter(created_by=request.user)

    context = {
        'form': form,
        'managed_users': managed_users,
        'is_super_admin': request.user.role == 'super_admin'
    }
    return render(request, 'my_team.html', context)


@login_required
def manage_projects_view(request):
    user = request.user
    if user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    if request.method == 'POST' and 'create_project' in request.POST:
        form = ClientProjectForm(request.POST, request.FILES)
        if form.is_valid():
            project = form.save(commit=False)
            project.created_by = user
            project.managed_by = user
            project.save()
            messages.success(request, f"Project '{project.name}' created.")
            return redirect('manage_projects')
    elif request.method == 'POST' and 'add_slot' in request.POST:
        slot_form = AdminSlotForm(request.POST)
        project_id = request.POST.get('project_id')

        if slot_form.is_valid() and project_id:
            project = get_object_or_404(ClientProject, id=project_id)
            slot_count = slot_form.cleaned_data['slot_count']
            for _ in range(slot_count):
                WorkEntry.objects.create(project=project, is_slot=True, user=None, folder_name="N/A", quantity=0)
            messages.success(request, f"{slot_count} new slot(s) added under project '{project.name}'.")
            return redirect('manage_projects')
    else:
        form = ClientProjectForm()

    if user.role == 'super_admin':
        projects = ClientProject.objects.prefetch_related('work_entries__category').all().order_by('-start_date')
    else:
        projects = ClientProject.objects.filter(managed_by=user).prefetch_related('work_entries__category').order_by('-start_date')

    slot_form = AdminSlotForm()

    context = {
        'form': form,
        'projects': projects,
        'slot_form': slot_form,
    }
    return render(request, 'manage_projects.html', context)


@login_required
def delete_project_view(request, project_id):
    user = request.user
    project = get_object_or_404(ClientProject, id=project_id)

    if user.role != 'super_admin' and project.managed_by != user:
        return render(request, 'unauthorized.html')

    project_name = project.name
    project.delete()
    messages.success(request, f"Project '{project_name}' has been deleted.")
    return redirect('manage_projects')


@login_required
def delete_user_view(request, user_id):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    user_to_delete = get_object_or_404(User, id=user_id)
    user_to_delete = get_object_or_404(User, id=user_id)

    if request.user.role == 'admin' and user_to_delete.created_by != request.user:
        messages.error(request, "You are not authorized to delete this user.")
        return redirect('my_team')

    if user_to_delete == request.user:
        messages.error(request, "You cannot delete your own account.")
        return redirect('my_team')

    user_name = user_to_delete.username
    user_to_delete.delete()
    messages.success(request, f"User '{user_name}' has been deleted successfully.")
    return redirect('my_team')


@login_required
def delete_work_entry_view(request, entry_id):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    work_entry = get_object_or_404(WorkEntry, id=entry_id)
    work_entry = get_object_or_404(WorkEntry, id=entry_id)

    if request.user.role == 'admin' and work_entry.project.managed_by != request.user:
        messages.error(request, "You are not authorized to delete this entry.")
        return redirect('manage_projects')

    work_entry.delete()
    messages.success(request, f"The work entry has been deleted successfully.")
    return redirect('manage_projects')


from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.utils import timezone
from django.conf import settings
from django.contrib import messages
from openpyxl import load_workbook
from openpyxl.styles import Border, Side
from copy import copy
from io import BytesIO
from num2words import num2words
import os

@login_required
def generate_invoice(request, project_id=None):
    user = request.user
    if user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    entries_qs = WorkEntry.objects.select_related('user', 'project', 'category')
    if user.role == 'admin':
        entries_qs = entries_qs.filter(project__managed_by=user)

    selected_project_id = project_id or request.GET.get('project')
    selected_user_id = request.GET.get('user')
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')

    if selected_project_id:
        entries_qs = entries_qs.filter(project_id=selected_project_id)
    if selected_user_id:
        entries_qs = entries_qs.filter(user_id=selected_user_id)
    if start_date:
        entries_qs = entries_qs.filter(date__date__gte=start_date)
    if end_date:
        entries_qs = entries_qs.filter(date__date__lte=end_date)

    work_entries = entries_qs.order_by('date')
    if not work_entries.exists():
        messages.warning(request, "No entries found to generate invoice.")
        return redirect('dashboard')

    template_path = os.path.join(settings.BASE_DIR, 'static', 'template', 'InvoiceTemplate.xlsx')
    try:
        wb = load_workbook(template_path)
        ws = wb.active
    except FileNotFoundError:
        return HttpResponse(f"Template not found at {template_path}", status=500)

    start_row = 14
    thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    footer_rows_data = []
    footer_start_row = -1
    for r_idx, row in enumerate(ws.iter_rows(min_row=start_row, max_row=ws.max_row), start=start_row):
        if isinstance(row[1].value, str) and "Total" in row[1].value:
            footer_start_row = r_idx
            break
    if footer_start_row != -1:
        for row in ws.iter_rows(min_row=footer_start_row, max_row=ws.max_row):
            footer_rows_data.append([copy(cell) for cell in row])
        ws.delete_rows(footer_start_row, ws.max_row - footer_start_row + 1)

    current_row = start_row
    total_quantity = 0
    total_amount = 0

    if len(work_entries) > 1:
        ws.insert_rows(start_row + 1, amount=len(work_entries) - 1)

    for entry in work_entries:
        unit_price = entry.category.rate if entry.category else 0
        quantity = entry.quantity or 0
        amount = quantity * unit_price

        ws.cell(row=current_row, column=2).value = entry.folder_name
        ws.cell(row=current_row, column=3).value = entry.category.name if entry.category else "N/A"
        local_date = timezone.localtime(entry.date)
        ws.cell(row=current_row, column=4).value = local_date.strftime("%b %d, %Y")
        ws.cell(row=current_row, column=5).value = quantity
        ws.cell(row=current_row, column=6).value = unit_price
        ws.cell(row=current_row, column=7).value = amount

        for col in range(2, 8):
            ws.cell(row=current_row, column=col).border = thin

        total_quantity += quantity
        total_amount += amount
        current_row += 1

    if footer_rows_data:
        for row_cells in footer_rows_data:
            ws.append([cell.value for cell in row_cells])
            new_row = ws.max_row
            for c_idx, cell in enumerate(row_cells, 1):
                if cell.has_style:
                    ws.cell(row=new_row, column=c_idx).style = copy(cell.style)

    for r_idx, row in enumerate(ws.iter_rows(min_row=current_row, values_only=True), start=current_row):
        if row[1] and "Total" in str(row[1]):
            ws.cell(row=r_idx, column=5).value = total_quantity
            ws.cell(row=r_idx, column=7).value = total_amount
        if row[1] and "In Words" in str(row[1]):
            first_entry = work_entries.first()
            currency = first_entry.category.currency if first_entry.category and hasattr(first_entry.category, 'currency') else 'USD'
            currency_map = {
                'USD': 'US Dollars',
                'EUR': 'Euros',
                'AUD': 'Australian Dollars',
                'INR': 'Rupees',
                'GBP': 'Pounds',
                'JPY': 'Yen',
            }
            currency_word = currency_map.get(currency.upper(), 'US Dollars')
            words = f"In Words: {num2words(total_amount, lang='en').title()} {currency_word} Only"
            ws.cell(row=r_idx, column=2).value = words

    excel_file_stream = BytesIO()
    wb.save(excel_file_stream)
    excel_file_stream.seek(0)

    project_for_invoice = None
    if selected_project_id:
        project_for_invoice = ClientProject.objects.get(id=selected_project_id)
    elif work_entries:
        project_for_invoice = work_entries.first().project

    if project_for_invoice:
        current_time = timezone.now()
        invoice = Invoice(
            project=project_for_invoice,
            project_name_snapshot=project_for_invoice.name,
            total_amount=total_amount,
            start_date=start_date if start_date else None,
            end_date=end_date if end_date else None
        )
        filename = f"Invoice_{project_for_invoice.name}_{current_time.year}_{current_time.month}.xlsx"
        invoice.invoice_file.save(filename, ContentFile(excel_file_stream.read()), save=True)
        excel_file_stream.seek(0)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response.write(excel_file_stream.read())

    return response

@login_required
def reports_view(request):
    user = request.user

    if user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    return render(request, 'reports.html')

@login_required
def invoice_reports_view(request):
    user = request.user

    if user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    if request.method == 'POST':
        invoice_ids = request.POST.getlist('invoice_ids')
        action = request.POST.get('bulk_action')

        if action == 'download':
            return bulk_download_invoices_view(request, invoice_ids)
        elif action == 'delete':
            return bulk_delete_invoices_view(request, invoice_ids)

    if user.role == 'super_admin':
        invoices = Invoice.objects.select_related('project').order_by('-generated_at')
        all_projects = ClientProject.objects.all()
    else:
        invoices = Invoice.objects.filter(project__managed_by=user).select_related('project').order_by('-generated_at')
        all_projects = ClientProject.objects.filter(managed_by=user)

    project_id = request.GET.get('project')
    month = request.GET.get('month')

    if project_id:
        invoices = invoices.filter(project_id=project_id)
    if month:
        try:
            year, month_num = month.split('-')
            invoices = invoices.filter(generated_at__year=year, generated_at__month=month_num)
        except ValueError:
            pass

    context = {
        'invoices': invoices,
        'all_projects': all_projects,
    }

    return render(request, 'invoice_reports.html', context)

@login_required
def delete_invoice_view(request, invoice_id):
    if request.user.role not in ['admin', 'super_admin']:
        return render(request, 'unauthorized.html')

    invoice = get_object_or_404(Invoice, id=invoice_id)

    if request.user.role == 'admin' and invoice.project.managed_by != request.user:
        messages.error(request, "You are not authorized to delete this report.")
        return redirect('reports')

    invoice.invoice_file.delete()
    invoice.delete()
    messages.success(request, "Invoice report deleted successfully.")
    return redirect('reports')

@login_required
def load_categories_view(request):
    project_id = request.GET.get('project_id')
    categories = []
    if project_id:
        categories = Category.objects.filter(project_id=project_id).order_by('name')

    return JsonResponse(list(categories.values('id', 'name')), safe=False)

@login_required
def bulk_download_invoices_view(request, invoice_ids):
    if not invoice_ids:
        messages.error(request, "No invoices selected for download.")
        return redirect('invoice_reports')

    invoices = Invoice.objects.filter(id__in=invoice_ids)
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for invoice in invoices:
            if invoice.invoice_file:
                zip_file.write(invoice.invoice_file.path, invoice.invoice_file.name)

    response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="invoices.zip"'
    return response

@login_required
def bulk_delete_invoices_view(request, invoice_ids):
    if not invoice_ids:
        messages.error(request, "No invoices selected for deletion.")
        return redirect('invoice_reports')

    invoices = Invoice.objects.filter(id__in=invoice_ids)
    for invoice in invoices:
        if request.user.role == 'admin' and invoice.project.managed_by != request.user:
            continue
        invoice.invoice_file.delete()
        invoice.delete()
    
    messages.success(request, "Selected invoices have been deleted.")
    return redirect('invoice_reports')

class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    serializer_class = RegisterSerializer
    permission_classes = [permissions.AllowAny]


class ProfileView(generics.RetrieveAPIView):
    serializer_class = UserSerializer
    permission_classes = [permissions.IsAuthenticated]
    def get_object(self):
        return self.request.user

class WorkEntryListCreateView(generics.ListCreateAPIView):
    serializer_class = WorkEntrySerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if user.role == 'super_admin':
            return WorkEntry.objects.all()
        if user.role == 'admin':
            return WorkEntry.objects.filter(project__managed_by=user)
        return WorkEntry.objects.filter(user=user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

class ExportWorkEntriesXLSXView(generics.GenericAPIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, *args, **kwargs):
        user = request.user
        if user.role not in ["admin", "super_admin"]:
            return Response({"detail": "Not authorized."}, status=403)

        return Response({"detail": "Export successful."})

class DashboardView(generics.ListAPIView):
    serializer_class = WorkDashboardSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if user.role == 'super_admin':
            queryset = WorkEntry.objects.all()
        elif user.role == 'admin':
            queryset = WorkEntry.objects.filter(project__managed_by=user)
        else:
            return WorkEntry.objects.none()
        return queryset

class PriceListCreateView(generics.ListCreateAPIView):
    queryset = Category.objects.all()
    serializer_class = CategorySerializer

    def get_queryset(self):
        user = self.request.user
        if user.role in ['super_admin', 'admin']:
            return Category.objects.filter(managed_by=user) if user.role == 'admin' else Category.objects.all()
        return Category.objects.none()

    def perform_create(self, serializer):
        serializer.save(managed_by=self.request.user)

class PriceDetailView(generics.RetrieveUpdateDestroyAPIView):
    queryset = Category.objects.all()
    serializer_class = CategorySerializer
    permission_classes = [permissions.IsAuthenticated]
    lookup_field = 'id'

    def get_queryset(self):
        user = self.request.user
        if user.role == 'super_admin':
            return Category.objects.all()
        if user.role == 'admin':
            return Category.objects.filter(managed_by=user)
        return Category.objects.none()

class ClientProjectListCreateView(generics.ListCreateAPIView):
    serializer_class = ClientProjectSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if user.role == 'super_admin':
            return ClientProject.objects.all()
        if user.role == 'admin':
            return ClientProject.objects.filter(managed_by=user)
        return ClientProject.objects.none()

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user, managed_by=self.request.user)

@login_required
def get_user_login_history(request, user_id):
    if request.user.role != 'super_admin':
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    user = get_object_or_404(User, id=user_id)
    login_history = user.login_history.all()[:5]  # Get last 5 logins
    
    history_data = []
    for login in login_history:
        # Convert to local timezone (Asia/Dhaka)
        local_datetime = timezone.localtime(login.login_datetime)
        history_data.append({
            'datetime': local_datetime.strftime('%Y-%m-%d %H:%M:%S'),
            'ip_address': login.ip_address,
            'device': f"{login.os} - {login.browser}",
            'device_type': login.device_type
        })
    
    return JsonResponse({'history': history_data})

class CustomLoginView(LoginView):
    template_name = 'login.html'
    form_class = AuthenticationForm

    def form_valid(self, form):
        response = super().form_valid(form)
        user = form.get_user()
        
        # Get client IP
        x_forwarded_for = self.request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip = x_forwarded_for.split(',')[0]
        else:
            ip = self.request.META.get('REMOTE_ADDR')
            
        # Parse user agent
        user_agent_string = self.request.META.get('HTTP_USER_AGENT', '')
        user_agent = user_agents.parse(user_agent_string)
        
        # Create login history entry
        UserLoginHistory.objects.create(
            user=user,
            ip_address=ip,
            user_agent=user_agent_string,
            device_type=user_agent.device.family,
            browser=user_agent.browser.family,
            os=user_agent.os.family
        )
        
        return response

class CustomLogoutView(LogoutView):
    template_name = 'registration/logged_out.html'
    next_page = 'home'

    def dispatch(self, request, *args, **kwargs):
        if request.method == 'GET':
            logout(request)
            return redirect(self.next_page)
        return super().dispatch(request, *args, **kwargs)

@login_required
def get_user_work_summary(request):
    """API endpoint to get filtered summary data for user's work entries dashboard"""
    if request.user.role != 'user':
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    # Base queryset
    work_entries = WorkEntry.objects.filter(user=request.user).select_related('project', 'category')
    
    # Filters
    query = request.GET.get('query')
    project_filter = request.GET.get('project')
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    
    # Build conditions
    search_conditions = Q()
    if query:
        search_conditions &= (Q(folder_name__icontains=query) | Q(category__name__icontains=query))
    if project_filter:
        search_conditions &= Q(project__name__icontains=project_filter)
    
    # ✅ Convert start_date and end_date safely
    if start_date_str:
        try:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            search_conditions &= Q(date__gte=start_date)
        except ValueError:
            pass
    
    if end_date_str:
        try:
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
            search_conditions &= Q(date__lte=end_date)
        except ValueError:
            pass
    
    # Apply filters
    filtered_entries = work_entries.filter(search_conditions)
    
    # Entries count
    total_entries = filtered_entries.count()
    
    # Total quantity (from filtered entries only)
    total_quantity = filtered_entries.aggregate(total=Sum('quantity'))['total'] or 0
    
    # Top project from filtered entries
    most_frequent_project = (
        filtered_entries.values('project__name')
        .annotate(count=Count('id'))
        .order_by('-count')
        .first()
    )
    top_project = most_frequent_project['project__name'] if most_frequent_project else 'No Projects'
    
    # Return JSON
    return JsonResponse({
        'total_entries_month': total_entries,        # filtered entries count
        'total_quantity_month': total_quantity,      # filtered total
        'most_frequent_project': top_project
    })

@login_required
def edit_user_role_view(request, user_id):
    if request.user.role != 'super_admin':
        return render(request, 'unauthorized.html')

    user_to_edit = get_object_or_404(User, id=user_id)
    
    if request.method == 'POST':
        form = UserRoleUpdateForm(request.POST, instance=user_to_edit)
        if form.is_valid():
            form.save()
            messages.success(request, f"Role updated successfully for user '{user_to_edit.username}'")
            return redirect('my_team')
    else:
        form = UserRoleUpdateForm(instance=user_to_edit)
    
    context = {
        'form': form,
        'user_to_edit': user_to_edit
    }
    return render(request, 'edit_user_role.html', context)

import zipfile

@login_required
def user_reports_view(request):
    user = request.user
    if user.role not in ['admin', 'super_admin', 'user']:
        return render(request, 'unauthorized.html')

    users_list = User.objects.select_related('created_by').order_by('username')

    if user.role == 'admin':
        users_list = users_list.filter(created_by=user)
    elif user.role == 'user':
        users_list = users_list.filter(id=user.id)

    search_query = request.GET.get('search', '')
    role_filter = request.GET.get('role', '')

    if search_query:
        users_list = users_list.filter(username__icontains=search_query)
    if role_filter:
        users_list = users_list.filter(role=role_filter)

    paginator = Paginator(users_list, 9)  # 9 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'page_obj': page_obj,
    }
    return render(request, 'user_reports.html', context)

@login_required
def get_user_report_view(request, user_id):
    try:
        print(f"[DEBUG] get_user_report_view called for user_id: {user_id}")
        print(f"[DEBUG] Requesting user role: {request.user.role}")
        
        if request.user.role not in ['admin', 'super_admin']:
            print("[DEBUG] Authorization failed - invalid role")
            return JsonResponse({'error': 'Unauthorized - Invalid role'}, status=403)

        user = get_object_or_404(User, id=user_id)
        print(f"[DEBUG] Found user: {user.username}")
        
        # Check permissions
        if request.user.role == 'admin' and user.created_by != request.user:
            print("[DEBUG] Authorization failed - admin permission check")
            return JsonResponse({'error': 'Unauthorized - Invalid permissions'}, status=403)

        # Get user's login history
        login_history = UserLoginHistory.objects.filter(user=user).order_by('-login_datetime')
        print(f"[DEBUG] Found {login_history.count()} login history records")
        
        # Get projects created by user
        projects_created = ClientProject.objects.filter(created_by=user).order_by('-start_date')
        print(f"[DEBUG] Found {projects_created.count()} projects created")
        
        # Get projects managed by user (if admin)
        projects_managed = ClientProject.objects.filter(managed_by=user).order_by('-start_date')
        print(f"[DEBUG] Found {projects_managed.count()} projects managed")
        
        # Get work entries with related data
        work_entries = WorkEntry.objects.filter(user=user).select_related(
            'project', 'category'
        ).order_by('-date')
        print(f"[DEBUG] Found {work_entries.count()} work entries")
        
        # Get work entries statistics
        work_stats = WorkEntry.objects.filter(user=user).aggregate(
            total_entries=models.Count('id'),
            total_quantity=models.Sum('quantity')
        )
        
        # Get project statistics
        project_stats = {}
        for entry in work_entries:
            if entry.project.id not in project_stats:
                project_stats[entry.project.id] = {
                    'name': entry.project.name,
                    'total_entries': 0,
                    'total_quantity': 0
                }
            project_stats[entry.project.id]['total_entries'] += 1
            project_stats[entry.project.id]['total_quantity'] += entry.quantity

        data = {
            'username': user.username,
            'role': user.get_role_display(),
            'created_by': user.created_by.username if user.created_by else None,
            'date_joined': user.date_joined.strftime('%Y-%m-%d'),
            
            # User summary
            'summary': {
                'total_work_entries': work_stats['total_entries'] or 0,
                'total_quantity': work_stats['total_quantity'] or 0,
                'projects_created_count': projects_created.count(),
                'projects_managed_count': projects_managed.count() if user.role in ['admin', 'super_admin'] else 0,
            },
            
            # Login history
            'login_history': [
                {
                    'datetime': timezone.localtime(lh.login_datetime).strftime('%Y-%m-%d %H:%M:%S'),
                    'ip_address': lh.ip_address,
                    'device': f"{lh.os} on {lh.browser}",
                    'browser': lh.browser,
                    'os': lh.os
                } for lh in login_history
            ],
            
            # Projects created by user
            'projects_created': [
                {
                    'name': p.name,
                    'start_date': p.start_date.strftime('%Y-%m-%d'),
                    'end_date': p.end_date.strftime('%Y-%m-%d') if p.end_date else 'Ongoing',
                    'managed_by': p.managed_by.username if p.managed_by else None,
                    'categories_count': p.categories.count()
                } for p in projects_created
            ],
            
            # Projects managed (if admin)
            'projects_managed': [
                {
                    'name': p.name,
                    'start_date': p.start_date.strftime('%Y-%m-%d'),
                    'end_date': p.end_date.strftime('%Y-%m-%d') if p.end_date else 'Ongoing',
                    'created_by': p.created_by.username,
                    'categories_count': p.categories.count()
                } for p in projects_managed
            ] if user.role in ['admin', 'super_admin'] else [],
            
            # Work entries
            'work_entries': [
                {
                    'date': we.date.strftime('%Y-%m-%d %H:%M'),
                    'project': we.project.name,
                    'category': we.category.name if we.category else 'N/A',
                    'folder_name': we.folder_name,
                    'quantity': we.quantity,
                    'rate': float(we.category.rate) if we.category else 0,
                    'currency': we.category.currency if we.category else 'USD'
                } for we in work_entries
            ],
            
            # Project-wise statistics
            'project_statistics': [
                {
                    'project_name': stats['name'],
                    'total_entries': stats['total_entries'],
                    'total_quantity': stats['total_quantity']
                } for project_id, stats in project_stats.items()
            ]
        }
        print(f"[DEBUG] Successfully prepared data for user {user.username}")
        return JsonResponse(data)
        
    except Exception as e:
        print(f"[DEBUG] Error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def export_user_report_view(request, user_id, format):
    if request.user.role not in ['admin', 'super_admin']:
        return HttpResponse("Unauthorized", status=403)

    user = get_object_or_404(User, id=user_id)

    if request.user.role == 'admin' and user.created_by != request.user:
        return HttpResponse("Unauthorized", status=403)

    # Fetch data (similar to get_user_report_view)
    login_history = UserLoginHistory.objects.filter(user=user).order_by('-login_datetime')[:10]
    projects_created = ClientProject.objects.filter(created_by=user).order_by('-start_date')
    work_entries = WorkEntry.objects.filter(user=user).select_related('project', 'category').order_by('-date')[:20]

    data = {
        'username': user.username,
        'role': user.get_role_display(),
        'created_by': user.created_by.username if user.created_by else None,
        'date_joined': user.date_joined.strftime('%Y-%m-%d'),
        'login_history': [
            {
                'datetime': timezone.localtime(lh.login_datetime).strftime('%Y-%m-%d %H:%M:%S'),
                'ip_address': lh.ip_address,
                'device': f"{lh.os} on {lh.browser}"
            } for lh in login_history
        ],
        'projects_created': [
            {
                'name': p.name,
                'start_date': p.start_date.strftime('%Y-%m-%d')
            } for p in projects_created
        ],
        'work_entries': [
            {
                'date': we.date.strftime('%Y-%m-%d'),
                'project': we.project.name,
                'category': we.category.name if we.category else 'N/A',
                'folder_name': we.folder_name,
                'quantity': we.quantity
            } for we in work_entries
        ],
    }

    if format == 'json':
        response = JsonResponse(data, content_type='application/json')
        response['Content-Disposition'] = f'attachment; filename="{user.username}_report.json"'
        return response
    elif format == 'docx':
        from docx import Document
        from docx.shared import Inches
        document = Document()
        document.add_heading(f'User Activity Report: {user.username}', 0)

        document.add_heading('User Info', level=1)
        document.add_paragraph(f"Role: {data['role']}")
        document.add_paragraph(f"Created By: {data['created_by']}")
        document.add_paragraph(f"Joined: {data['date_joined']}")

        document.add_heading('Login History', level=1)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Timestamp'
        hdr_cells[1].text = 'IP Address'
        hdr_cells[2].text = 'Device'
        for login in data['login_history']:
            row_cells = table.add_row().cells
            row_cells[0].text = login['datetime']
            row_cells[1].text = login['ip_address']
            row_cells[2].text = login['device']

        document.add_heading('Projects Created', level=1)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Start Date'
        for project in data['projects_created']:
            row_cells = table.add_row().cells
            row_cells[0].text = project['name']
            row_cells[1].text = project['start_date']

        # Create a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{user.username}_report.docx"'
        document.save(response)
        return response

    return HttpResponse("Invalid format.", status=400)
